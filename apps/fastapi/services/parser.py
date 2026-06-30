"""
CV/document parser — turns an uploaded PDF or DOCX into plain text. PyMuPDF (fitz) handles
the common case fast; pdfplumber is used as a fallback specifically for table-heavy pages
where PyMuPDF's text extraction tends to garble column layout. python-docx covers Word
files. Unstructured.io is deliberately not used (system-dep issues in the dev container —
see ADR-008). Two guards run after extraction: a minimum-length check (cv_unreadable —
catches blank/scanned PDFs) and a token-budget truncation so the LLM extraction step in
1.3 never receives more than CV_MAX_TOKENS, prioritising the front of the document where
expertise/skills sections usually live.
"""
from __future__ import annotations

import io
import re

import fitz  # PyMuPDF
import pdfplumber
from core.errors import ERRORS, APIError
from core.logging import get_logger
from docx import Document

from core.config import settings

logger = get_logger(__name__)

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPES = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
)

# Rough chars-per-token heuristic for the pre-LLM truncation guard (no tokenizer call needed
# at parse time; the LLM call itself enforces the real limit, this just avoids sending an
# obviously oversized payload).
_CHARS_PER_TOKEN_ESTIMATE = 4


def parse_cv(file_bytes: bytes, content_type: str) -> str:
    """
    Returns extracted, truncated, whitespace-normalised plain text from a PDF or DOCX file.
    Raises APIError with a stable code on any failure — never lets a raw exception escape.
    """
    if content_type == PDF_CONTENT_TYPE:
        text = _parse_pdf(file_bytes)
    elif content_type in DOCX_CONTENT_TYPES:
        text = _parse_docx(file_bytes)
    else:
        raise APIError(
            "unsupported_file_type", ERRORS["unsupported_file_type"], status_code=400
        )

    text = _normalise_whitespace(text)
    _guard_min_length(text)
    return _truncate_to_token_budget(text)


def _parse_pdf(file_bytes: bytes) -> str:
    """
    Primary path: PyMuPDF, fast and good on simple single-column CVs. If the result looks
    suspiciously short relative to the page count (a sign of a table-heavy or oddly laid
    out CV), fall back to pdfplumber, which handles tables and multi-column text better.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = doc.page_count
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
    except Exception as exc:
        logger.warning("PyMuPDF failed to open PDF: %s", exc)
        text, page_count = "", 1

    # Heuristic: fewer than ~40 chars/page strongly suggests a scanned image or a layout
    # PyMuPDF mishandled (e.g. heavy tables) — try pdfplumber before giving up.
    if page_count and len(text.strip()) < 40 * page_count:
        fallback = _parse_pdf_with_pdfplumber(file_bytes)
        if len(fallback.strip()) > len(text.strip()):
            return fallback

    return text


def _parse_pdf_with_pdfplumber(file_bytes: bytes) -> str:
    try:
        chunks: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                chunks.append(page_text)
                for table in page.extract_tables():
                    for row in table:
                        chunks.append(" | ".join(cell or "" for cell in row))
        return "\n".join(chunks)
    except Exception as exc:
        logger.warning("pdfplumber fallback failed: %s", exc)
        return ""


def _parse_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.warning("python-docx failed to open file: %s", exc)
        return ""


def _normalise_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _guard_min_length(text: str) -> None:
    """First failure guard (1.2b): blank/scanned/unreadable PDFs surface a typed error,
    never a 500 — the frontend shows a 'paste your CV text' textarea on this code."""
    if len(text) < settings.CV_MIN_CHARS:
        raise APIError("cv_unreadable", ERRORS["cv_unreadable"], status_code=422)


def _truncate_to_token_budget(text: str) -> str:
    """Second guard (1.2b): keep the document under CV_MAX_TOKENS before the LLM call.
    Prioritises the start of the document, where expertise/skills/summary sections
    conventionally appear on a CV, rather than truncating arbitrarily from the end only
    if the budget allows — for very long CVs we keep the head and a smaller tail slice
    so trailing publication lists aren't entirely lost."""
    max_chars = settings.CV_MAX_TOKENS * _CHARS_PER_TOKEN_ESTIMATE
    if len(text) <= max_chars:
        return text

    head_budget = int(max_chars * 0.85)
    tail_budget = max_chars - head_budget
    head = text[:head_budget]
    tail = text[-tail_budget:] if tail_budget > 0 else ""
    logger.info(
        "CV text truncated from %d to %d chars (head+tail budget).", len(text), max_chars
    )
    return f"{head}\n\n[... truncated ...]\n\n{tail}"
